from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import random
import zipfile
import tempfile
import shutil

app = Flask(__name__)

# ✅ Fixes placeholders even when split across runs, preserves header/footer
def replace_all_text(doc, replacements):
    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            replaced = False
            for key, val in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, val.upper())
                    replaced = True
            if replaced and para.runs:
                for run in para.runs:
                    run.text = ''
                para.runs[0].text = full_text

    def replace_in_table_cells(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

    replace_in_paragraphs(doc.paragraphs)
    replace_in_table_cells(doc.tables)

    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_table_cells(section.header.tables)
        replace_in_paragraphs(section.footer.paragraphs)
        replace_in_table_cells(section.footer.tables)

# ✅ Randomize values in test columns
def update_measured_values(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 14:
                sr_no = cells[0].text.strip()
                if sr_no in ["1", "2"]:
                    value = f"{random.randint(54, 94) if sr_no == '1' else random.randint(112, 194)}µA"
                    for i in range(11, 14):
                        cell = cells[i]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        run = para.add_run(value)
                        run.font.name = 'Cambria'
                        run.font.size = Pt(11)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        excel_file = request.files["excel"]
        word_file = request.files["word"]

        temp_dir = tempfile.mkdtemp()
        try:
            excel_path = os.path.join(temp_dir, "input.xlsx")
            word_path = os.path.join(temp_dir, "template.docx")
            excel_file.save(excel_path)
            word_file.save(word_path)
            print(f"📥 Files saved: {excel_path}, {word_path}")

            df = pd.read_excel(excel_path)
            df.columns = df.columns.str.strip()

            output_dir = os.path.join(temp_dir, "output")
            os.makedirs(output_dir, exist_ok=True)

            for _, row in df.iterrows():
                clr = str(row.get('CLR', '')).strip()
                ulr = str(row.get('ULR', '')).strip()
                doc = Document(word_path)

                replacements = {
                    "<Equipment name>": str(row.get("Equipment name", "")),
                    "<MK>": str(row.get("MK", "")),
                    "<MO>": str(row.get("MO", "")) if pd.notna(row.get("MO")) else "",
                    "<SN>": str(row.get("SN", "")),
                    "<ID>": str(row.get("ID", "")),
                    "<DEPT>": str(row.get("DEPT", "")),
                    "<D Date>": str(row.get("D Date", "")),
                    "<E date>": str(row.get("E date", "")),
                    "<CLR>": clr,
                    "<ULR>": ulr,
                    "<TEM>": str(row.get("TEM", "")),
                    "<HUM>": str(row.get("HUM", "")),
                }

                print(f"🧩 Replacing placeholders for CLR={clr}")
                replace_all_text(doc, replacements)
                update_measured_values(doc)

                output_filename = f"EST - {clr}.docx"
                output_path = os.path.join(output_dir, output_filename)
                doc.save(output_path)
                print(f"✅ Saved: {output_filename}")

            zip_path = os.path.join(temp_dir, "result.zip")
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                for file in os.listdir(output_dir):
                    zipf.write(os.path.join(output_dir, file), arcname=file)

            stable_zip_path = os.path.join(tempfile.gettempdir(), f"result_{random.randint(1000,9999)}.zip")
            shutil.move(zip_path, stable_zip_path)

            response = send_file(
                stable_zip_path,
                as_attachment=True,
                download_name="documents.zip",
                mimetype="application/zip"
            )
            response.call_on_close(lambda: os.remove(stable_zip_path))
            return response

        except Exception as e:
            print(f"❌ Error: {e}")
            return "An error occurred while generating documents", 500

        finally:
            try:
                shutil.rmtree(temp_dir)
                print(f"🗑️ Temp cleaned: {temp_dir}")
            except Exception as e:
                print(f"⚠️ Cleanup failed: {e}")

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
